import json
import os
import re
import time
import uuid
from src.core.config_manager import ConfigManager
from src.core.core_task import BackgroundTask, TaskState
from src.core.device_manager import DeviceManager
from src.core.kb_manager import KBManager, DatabaseManager
from src.core.mcp_manager import MCPManager
from src.core.models_registry import get_model_conf, resolve_auto_model


class ChatGenerationTask(BackgroundTask):
    """
    Background task for executing local/remote LLM interactions, Vector Retrieval,
    and Multi-Agent tool processing within the Core Task Framework.
    """

    # 首次重排失败弹一次警告，之后静默降级，避免每次问答刷屏
    _rerank_warned = False

    def cancel(self):
        super().cancel()
        if hasattr(self, 'main_llm') and self.main_llm: self.main_llm.cancel()
        if hasattr(self, 'trans_llm') and self.trans_llm: self.trans_llm.cancel()
        if hasattr(self, 'vision_llm') and self.vision_llm: self.vision_llm.cancel()

    def _init_llms(self):
        from src.core.llm_impl import OpenAICompatibleLLM
        if self.main_config and not getattr(self, 'main_llm', None):
            cfg = self.main_config.copy()
            if "tools" not in cfg:
                cfg["tools"] = []

            if "extra_params" not in cfg:
                cfg["extra_params"] = {}
            cfg["extra_params"]["timeout"] = 600.0
            cfg["timeout"] = 600.0
            self.main_llm = OpenAICompatibleLLM(cfg)

        if self.requires_translation and self.trans_config and not getattr(self, 'trans_llm', None):
            self.trans_llm = OpenAICompatibleLLM(self.trans_config)

    def _emit_token(self, token: str):
        self.update_progress(-1, token)

    def _emit_error(self, msg: str):
        raise RuntimeError(msg)

    def _emit_translated(self, text: str):
        self._emit_state(TaskState.PROCESSING, -1, "", payload={"event": "translated", "text": text})

    def _execute(self):
        from src.core.llm_impl import OpenAICompatibleLLM, get_cached_translation

        self.send_log("INFO", f"Chat task started. KB_ID: {self.kwargs.get('kb_id')}")
        time.sleep(0.1)

        self.main_config = self.kwargs.get('main_config')
        self.trans_config = self.kwargs.get('trans_config')
        self.messages = self.kwargs.get('messages', [])
        self.kb_id = self.kwargs.get('kb_id')
        if self.kb_id == "none": self.kb_id = None

        current_external_files = self.kwargs.get('external_files', [])
        all_external_files = []

        # 1. 遍历历史获取上下文遗留文件
        for m in self.messages:
            if m.get('external_files'):
                for f in m['external_files']:
                    if f not in all_external_files:
                        all_external_files.append(f)

        # 2. 合并当前上传文件
        for f in current_external_files:
            if f not in all_external_files:
                all_external_files.append(f)

        self.external_context = []

        if all_external_files:
            self.send_log("INFO", f"Loading {len(all_external_files)} attached file(s) into memory context...")
            self._emit_token(
                f"<div class='status-msg' style='color:#05B8CC; margin-bottom:4px;'>📄 Loading {len(all_external_files)} attached file(s) into memory...</div>\n\n")
            time.sleep(0.05)

            import tempfile, hashlib, os
            cache_dir = os.path.join(tempfile.gettempdir(), "scholar_navis_cache")
            os.makedirs(cache_dir, exist_ok=True)

            for info in all_external_files:
                if self.is_cancelled(): break
                path = info.get('path', '')
                f_name = info.get('name', 'Unknown')
                content = info.get('content', None)
                ext = f_name.lower()

                if content is not None:
                    self.external_context.append(
                        {"path": path, "name": f_name, "page": info.get('page', 1), "content": content})
                    continue

                if os.path.exists(path):
                    file_stat = os.stat(path)
                    hash_key = hashlib.md5(f"{path}_{file_stat.st_mtime}_{file_stat.st_size}".encode()).hexdigest()
                    cache_file = os.path.join(cache_dir, f"{hash_key}.json")

                    # 击中缓存，直接加载，实现秒进
                    if os.path.exists(cache_file):
                        try:
                            with open(cache_file, 'r', encoding='utf-8') as cf:
                                cached_data = json.load(cf)
                            self.external_context.extend(cached_data)
                            continue
                        except:
                            pass

                    try:
                        chunks = []
                        if ext.endswith('.pdf'):
                            import pymupdf4llm
                            md_chunks = pymupdf4llm.to_markdown(path, page_chunks=True)
                            for chunk in md_chunks:
                                text = chunk.get("text", "").strip()
                                if len(text) > 10:
                                    chunks.append({
                                        "path": path, "name": f_name, "page": chunk.get("metadata", {}).get("page", 1),
                                        "content": text
                                    })
                        elif ext.endswith('.docx'):
                            import docx
                            doc = docx.Document(path)
                            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                            if len(text) > 10:
                                chunks.append({"path": path, "name": f_name, "page": 1, "content": text})
                        elif ext.endswith('.doc'):
                            self.send_log("WARNING", f"Legacy .doc format skipped: {f_name}")
                        elif ext.endswith(('.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp')):
                            self.send_log("INFO", f"Image upload is currently paused. Skipping image: {f_name}")
                        else:
                            import chardet
                            with open(path, 'rb') as f:
                                raw_data = f.read()
                                detected = chardet.detect(raw_data)
                                encoding = detected['encoding'] if detected['encoding'] else 'utf-8'
                                text = raw_data.decode(encoding, errors='replace').strip()
                            if text:
                                chunks.append({"path": path, "name": f_name, "page": 1, "content": text})

                        if chunks:
                            self.external_context.extend(chunks)
                            with open(cache_file, 'w', encoding='utf-8') as cf:
                                json.dump(chunks, cf, ensure_ascii=False)
                    except Exception as e:
                        self.send_log("ERROR", f"Failed to parse {f_name}: {e}")

            self._emit_token("[CLEAR_SEARCH]")

        if self.kb_id:
            from src.core.models_registry import ModelManager
            ready, missing_label, missing_id, m_type = ModelManager().verify_chat_models(self.kb_id)
            if not ready:
                self._emit_error(json.dumps({
                    "title": "Model Missing - Action Blocked",
                    "body": f"Required offline model is not installed:\n• {missing_label}\n\nPlease go to [Global Settings] and click 'Save' to download required models."
                }))
                return

        original_user_query = self.messages[-1].get('display_text', self.messages[-1].get('content', ''))

        try:
            from src.core.lang_detect import detect_primary_language
            is_english = detect_primary_language(original_user_query) == 'en'

            if not is_english and self.trans_config is None:
                self.send_log("WARNING",
                              "Non-English input detected, but translation model is not enabled. The core model may not perfectly handle this language.")
                self.requires_translation = False
            else:
                self.requires_translation = (not is_english)

        except Exception as e:
            self.logger.warning(f"Language detection failed in background: {e}")
            self.requires_translation = False

        self.use_academic_agent = self.kwargs.get('use_academic_agent', True)
        self.academic_tags = self.kwargs.get('academic_tags', [])


        self.use_external_tools = self.kwargs.get('use_external_tools', False)
        self.external_tool_names = self.kwargs.get('external_tool_names',
                                                   [])

        # 深度研究开关：由 UI 传入（kwargs 优先），否则回退到全局配置
        self.deep_mode = self.kwargs.get('deep_mode', None)
        if self.deep_mode is None:
            self.deep_mode = bool(self.config.user_settings.get("agent_deep_mode", False))

        self.db = DatabaseManager()
        self.kb_manager = KBManager()
        self.config = ConfigManager()
        self.full_response_cache = ""

        # New conversation: clear the previous round's Provenance evidence chain
        # to avoid unbounded cross-session accumulation (the collector is a
        # process-level singleton).
        try:
            from src.core.provenance import get_collector
            get_collector().clear()
        except Exception as e:
            self.logger.warning(f"Failed to clear provenance collector: {e}")

        # 后台预加载 Reranker，避免首次问答时主线程阻塞（失败会惰性重试）
        if self.config.user_settings.get("rerank_auto_load", True):
            try:
                import threading as _t
                _t.Thread(target=self._preload_reranker, daemon=True).start()
            except Exception as e:
                self.logger.warning(f"Failed to spawn reranker preload thread: {e}")

        self.main_llm = None
        self.trans_llm = None
        self.vision_llm = None

        self._init_llms()
        # 新一轮生成开始：复位各 LLM 实例的取消标志（若实例复用）。
        # 放在任务入口而非 AgentRuntime.run 内，避免 deep 模式多子 Agent 共享
        # 同一 LLM 时，并发 reset 与用户取消产生竞态。
        for llm in (self.main_llm, self.trans_llm, self.vision_llm):
            if llm is not None and hasattr(llm, "reset"):
                try:
                    llm.reset()
                except Exception as e:
                    self.logger.warning(f"Failed to reset LLM cancel state: {e}")

        original_user_query = self.messages[-1].get('display_text', self.messages[-1].get('content', ''))
        search_query = original_user_query
        domain = "General Academic"
        context_str = ""
        sources_map = {}

        # Phase 1: Query Extraction & Translation (Cache Accelerated)
        if self.requires_translation:
            self.send_log("INFO", f"Translating query: {original_user_query[:20]}...")
            self._emit_token(
                "<div class='status-msg' style='color:#05B8CC; margin-bottom:4px;'>🌐 Translating your query to academic English for precise retrieval...</div>\n\n")
            try:
                trans_kwargs = {
                    "is_translation": True,
                    "stream": False
                }
                search_query = get_cached_translation(original_user_query, "to_en", self.trans_llm, **trans_kwargs)
                self._emit_translated(search_query)
            except Exception as e:
                self._emit_error(f"Translation model request failed. Details: {e}")

        # Phase 2: Vector Retrieval & Reranking (Local KB)
        if self.kb_id:
            self.send_log("INFO", "Initiating local Vector RAG retrieval...")
            self._emit_token("[CLEAR_SEARCH]")
            self._emit_token(
                "<div class='status-msg' style='color:#05B8CC; margin-bottom:4px;'>📚 Searching local knowledge base and reranking documents...</div>\n\n")
            time.sleep(0.05)

            kb_info = self.kb_manager.get_kb_by_id(self.kb_id)
            if kb_info and kb_info.get('doc_count', 0) == 0:
                self.logger.warning(f"Knowledge Base '{kb_info.get('name')}' is empty. Skipping vector retrieval.")
            elif kb_info:
                self._emit_token(
                    "<div class='status-msg' style='color:#05B8CC; margin-bottom:4px;'>Loading local vector model and retrieving literature...</div>\n\n")
                domain = kb_info.get('domain', 'General Academic')
                model_id = kb_info.get('model_id', 'embed_auto')

                user_pref = self.config.user_settings.get("inference_device", "Auto")
                target_device = DeviceManager().parse_device_string(user_pref)

                conf = get_model_conf(model_id, "embedding")
                if not conf or conf.get('is_auto'):
                    from src.task.kb_tasks import _worker_load_model
                    real_id = resolve_auto_model("embedding", target_device)
                    conf = get_model_conf(real_id, "embedding")

                try:
                    from src.task.kb_tasks import _worker_load_model
                    embed_fn = _worker_load_model(self.kb_id, self.config)
                    if not self.db.switch_kb(self.kb_id, embedding_function=embed_fn):
                        self._emit_error(f"Failed to switch to Knowledge Base: {self.kb_id}")
                except Exception as e:
                    self._emit_error(f"Critical Model Error: {str(e)}")

                history_context = ""
                if len(self.messages) >= 3:
                    prev_assistant = self.messages[-2]['content'][:100]
                    history_context = f" (Context: {prev_assistant})"

                expanded_queries = [
                    search_query,
                    f"{search_query}{history_context}",
                    f"{domain} context: {search_query} research details"
                ]

                candidate_docs = []
                seen_contents = set()

                for eq in expanded_queries:
                    raw_results = self.db.query(eq, n_results=20)
                    if raw_results and raw_results.get('documents') and raw_results['documents'][0]:
                        docs = raw_results['documents'][0]
                        metas = raw_results['metadatas'][0]
                        distances = raw_results.get('distances', [[0] * len(docs)])[0]

                        for i, doc_text in enumerate(docs):
                            clean_text = doc_text.strip()
                            if clean_text not in seen_contents and len(clean_text) > 20:
                                seen_contents.add(clean_text)
                                candidate_docs.append({
                                    "content": clean_text,
                                    "metadata": metas[i],
                                    "v_dist": distances[i]
                                })

                if candidate_docs:
                    candidate_docs = sorted(candidate_docs, key=lambda x: x.get('v_dist', 0))[:40]
                    final_docs = self._process_rerank(search_query, candidate_docs, domain)
                    if final_docs is None:
                        final_docs = candidate_docs[:10]

                    current_ref_id = 1
                    for doc in final_docs:
                        sources_map[current_ref_id] = {
                            "path": doc['metadata'].get('file_path', ''),
                            "page": doc['metadata'].get('page', 1),
                            "name": doc['metadata'].get('source', 'Local DB'),
                            "search_text": doc['content'][:100]
                        }
                        context_str += (
                            f"--- [Document {current_ref_id}] ---\n"
                            f"Source: {doc['metadata'].get('source', 'Local')}\n"
                            f"Content: {doc['content']}\n\n"
                        )
                        current_ref_id += 1

        if not context_str.strip():
            context_str = "No local database documents provided."

        external_chunks = self.external_context or []
        images = [c for c in external_chunks if c.get("type") == "image" or str(c.get("path", "")).lower().endswith(
            ('.png', '.jpg', '.jpeg', '.webp'))]
        docs = [c for c in external_chunks if c not in images]

        llm_content = []

        # Phase 3: External Attachments Integration(文件打分环节)
        if docs:
            self.send_log("INFO", f"Detected {len(docs)} uploaded document chunks. Starting Reranker scoring...")
            self._emit_token(
                "<div class='status-msg' style='color:#05B8CC; margin-bottom:4px;'>Filtering and reranking attached documents...</div>\n\n")
            cand_docs = [{"content": d.get("content", ""),
                          "metadata": {"name": d.get("name", "Unknown"), "page": d.get("page", 1)}} for d in docs]

            if len(cand_docs) > 5:
                reranked_docs = self._process_rerank(search_query, cand_docs, "General")
                if reranked_docs is not None:
                    self.send_log("INFO",
                                  f"Reranker finished: Reduced {len(cand_docs)} chunks to top {len(reranked_docs)} most relevant segments.")
                    cand_docs = reranked_docs
                else:
                    self.send_log("WARNING", "Reranker failed for files, falling back to top-k selection.")
                    cand_docs = cand_docs[:8]
            else:
                self.send_log("INFO",
                              f"Small attachment size ({len(cand_docs)} chunks), skipping rerank and using all content.")

            for doc in cand_docs:
                f_name = doc["metadata"]["name"]
                page = doc["metadata"]["page"]
                context_str += (
                    f"--- [User Attached File: {f_name} (Page {page})] ---\n"
                    f"Content: {doc['content']}\n\n"
                )

        if images:
            vision_model_name = self.main_config.get("vision_model_name", "auto")
            main_model_name = self.main_config.get("model_name", "").lower()

            vision_keywords = ['image', 'vl', 'vision', 'llava', 'pixtral', 'gpt-4o', 'gpt-4-turbo', 'gemini-1.5',
                               'gemini-2.0', 'claude-3', 'qwen-vl']
            main_supports_vision = any(kw in main_model_name for kw in vision_keywords)
            if "deepseek" in main_model_name:
                main_supports_vision = False

            need_pre_caption = False
            active_vision_model = None

            if vision_model_name != "auto":
                need_pre_caption = True
                active_vision_model = vision_model_name
            elif not main_supports_vision:
                need_pre_caption = True
                active_vision_model = main_model_name

            if need_pre_caption:
                self._emit_token(
                    "<div class='status-msg' style='color:#05B8CC; margin-bottom:4px;'>Extracting image contexts via vision model...</div>\n\n")
                try:
                    vision_cfg = self.main_config.copy()
                    vision_cfg["model_name"] = active_vision_model
                    vision_cfg.pop("tools", None)
                    self.vision_llm = OpenAICompatibleLLM(vision_cfg)

                    image_descriptions = []
                    for img in images:
                        if self.is_cancelled(): break

                        img_data = img.get("base64_url") or img.get("content")
                        if not img_data.startswith("data:image"):
                            ext = str(img.get("path", ".jpeg")).split('.')[-1]
                            img_data = f"data:image/{ext};base64,{img_data}"

                        vision_prompt = [{"role": "user", "content": [
                            {"type": "text",
                             "text": "Please deeply analyze this image, extract all text (OCR), describe the charts/data, and detail its core contents. Output in pure text."},
                            {"type": "image_url", "image_url": {"url": img_data}}
                        ]}]

                        desc_res = self.vision_llm.chat(vision_prompt)
                        desc_content = desc_res.get('content', '') if isinstance(desc_res, dict) else str(desc_res)

                        image_descriptions.append(
                            f"[Image: {img.get('name', 'Unknown')}] Description:\n{desc_content}")

                    if image_descriptions:
                        llm_content.append({"type": "text",
                                            "text": "The user uploaded images. Here are their detailed textual descriptions analyzed by the vision model:\n" + "\n".join(
                                                image_descriptions)})
                except Exception as e:
                    self.logger.warning(f"Vision pre-captioning failed: {e}")
                    self._emit_token(
                        "<div style='color:#e6a23c;'>⚠️ Image parsing failed. The current model configuration might not support vision. Images will be ignored.</div><br>")
                    llm_content.append({"type": "text",
                                        "text": f"[System Warning: User uploaded an image, but the vision parser failed to read it.]"})
                finally:
                    self.vision_llm = None
            else:
                self.logger.info(f"Mounting images natively for vision-capable main model: [{main_model_name}]")
                for img in images:
                    img_data = img.get("base64_url") or img.get("content")
                    if img_data:
                        if not img_data.startswith("data:image"):
                            ext = str(img.get("path", ".jpeg")).split('.')[-1]
                            img_data = f"data:image/{ext};base64,{img_data}"
                        llm_content.append({
                            "type": "image_url",
                            "image_url": {"url": img_data}
                        })

        llm_content.append({"type": "text", "text": f"User Query:\n{search_query}"})
        self._emit_token("[CLEAR_SEARCH]")

        # Phase 5: Agentic Generation (Modern Agent Runtime)
        self._emit_token("[START_LLM_NETWORK]")

        mcp_mgr = MCPManager.get_instance()
        from src.core.skill_manager import SkillManager
        from src.core.agent.skill_registry import SkillRegistry
        from src.core.agent.planner import IntentPlanner
        from src.core.agent.runtime import AgentRuntime

        skill_mgr = SkillManager.get_instance()
        registry = SkillRegistry(skill_mgr).build()
        planner = IntentPlanner(registry)

        raw_tools = []
        # 1. 内部学术 Agent (enabled Skills by user tags)
        if self.use_academic_agent:
            raw_academic = skill_mgr.get_academic_schemas(self.academic_tags)
            if raw_academic:
                raw_tools.extend(raw_academic)

        # 2. 外部工具组合 (external Skills + remote MCP by user names)
        if self.use_external_tools:
            ext_skills = skill_mgr.get_external_schemas(self.external_tool_names)
            if ext_skills:
                raw_tools.extend(ext_skills)

            for schema in mcp_mgr.tool_schemas.values():
                server_name = schema.get("server", "Unknown Server")
                if not self.external_tool_names or server_name in self.external_tool_names:
                    raw_tools.append({
                        "type": schema.get("type", "function"),
                        "function": schema.get("function", {})
                    })

        # --- Modern AGENT tool exposure (LLM decides via native function calling) ---
        # The Planner NEVER strips tools with keyword matching. All user-enabled
        # tools are exposed so the main LLM keeps full agency; an optional
        # semantic-focus hint is injected only as guidance.
        focus_reminder = ""
        if raw_tools:
            self.send_log(
                "INFO",
                f"Enabled tool pool: {len(raw_tools)} tools (Skills + MCP). "
                "Exposing all of them for native LLM function calling...",
            )
            self._emit_token(
                "<div class='status-msg' style='color:#05B8CC; margin-bottom:4px;'>"
                "Analyzing query intent to guide tool selection...</div>\n\n"
            )
            time.sleep(0.05)

            # Full tool set -> full agency for the LLM.
            combined_tools = list(raw_tools)

            # Optional semantic focus hint (suggestion only, never a filter).
            use_semantic_focus = self.config.user_settings.get("agent_semantic_focus", True)
            if use_semantic_focus:
                try:
                    plan = planner.semantic_focus(search_query, self.main_llm)
                    focus_reminder = planner.build_focus_reminder(plan)
                    if focus_reminder:
                        self.send_log(
                            "INFO",
                            f"Semantic focus hint: {', '.join(plan.focus_hint)} "
                            f"(intent='{plan.intent}').",
                        )
                except Exception as e:
                    self.logger.warning(f"Semantic focus failed, continuing with full tool set: {e}")

            self._emit_token("[CLEAR_SEARCH]")
            final_tool_names = [t.get("function", {}).get("name", "Unknown") for t in combined_tools]
            self.send_log(
                "INFO",
                f"Exposing {len(final_tool_names)} tools for LLM selection: {', '.join(final_tool_names)}",
            )
        else:
            combined_tools = []

        # Always expose the image generator and chart modifier so drawing requests
        # are never blocked and previously drawn charts can be re-rendered.
        from src.core.agent.runtime import _ALWAYS_TOOLS
        combined_tools.append(dict(_ALWAYS_TOOLS)["generate_image"])
        combined_tools.append(dict(_ALWAYS_TOOLS)["modify_chart"])
        combined_tools.append(dict(_ALWAYS_TOOLS)["propose_plot_plan"])

        if combined_tools:
            self._emit_token("<mcp_process>⚙️ Query intent analyzed — the model selects tools natively...</mcp_process>")
            tool_names = [t.get("function", {}).get("name", "Unknown") for t in combined_tools]
            # 当绘图工具可用时，追加一条强约束，避免 reasoning 模型把 chart 参数
            # 以纯文本形式“泄漏”到最终答案里，而不是真正调用 plot_chart。
            plot_guard = ""
            if any(name == "plot_chart" for name in tool_names):
                plot_guard = (
                    "\n### DATA VISUALIZATION RULE (plot_chart / propose_plot_plan):\n"
                    "When the user requests any chart or plot (bubble, bar, scatter, volcano, heatmap, "
                    "GO enrichment, volcano plot, etc.), you MUST call the plot_chart tool to render the "
                    "figure. NEVER reply with chart parameters, the data table, or a chart specification "
                    "as plain text.\n"
                    "IMPORTANT: If the user asks to visualize/plot data but has NOT clearly specified the "
                    "chart type, the x/y columns, the title, or styling, call the propose_plot_plan tool "
                    "FIRST to show a confirmation card. Only call plot_chart after the user confirms the plan.\n"
                    "ENRICHMENT DOTPLOT REFERENCE LAYOUT (KEGG / GO / GSEA / Reactome bubble plots): "
                    "chart_type='bubble', x=Gene Ratio plotted HORIZONTALLY at the bottom, y=Pathway/Term "
                    "name on the left ordered by Gene Ratio DESCENDING (largest ratio on TOP), size=Gene "
                    "Count, color=FDR (BH-corrected p-value) with a blue-to-red continuous gradient; the "
                    "right-side legend has a vertical color bar labeled 'FDR' plus a 'Count' size legend "
                    "with discrete reference dots. Do NOT coord_flip; do NOT use -log10(FDR) for the "
                    "color mapping (use the raw FDR column directly so the gradient matches the reference).\n"
                    "If native function calling is unavailable, output exactly this JSON block and nothing else:\n"
                    "```json {\"name\": \"plot_chart\", \"arguments\": {\"chart_type\": \"bubble\", \"data\": \"[...]\", \"x\": \"...\", \"y\": \"...\"}} ```\n"
                )
            dynamic_tool_prompt = (
                f"### CRITICAL TOOL UTILIZATION RULE:\n"
                f"You have the following tools available for this query: {', '.join(tool_names)}.\n"
                f"Read the user's prompt carefully and USE the native function-calling API to invoke the "
                f"tool(s) you need. If the user asks for multi-dimensional data (e.g., metadata AND protein "
                f"interactions), you MUST use multiple tools to fulfill ALL parts of the request. DO NOT skip "
                f"required tools. DO NOT answer partially.\n\n"
                f"{focus_reminder}{plot_guard}"
            )
        else:
            dynamic_tool_prompt = ""

        system_prompt = (
            f"You are a Senior Research Scientist specializing in {domain}. "
            "Your goal is to provide high-density, evidence-based academic responses.\n\n"
            f"{dynamic_tool_prompt}\n\n"
            "### TOOL USE PROTOCOL (STRICT):\n"
            "1. CRITICAL FOR CITATIONS: If the user's prompt asks for literature, references, citations, or a review, you MUST explicitly invoke academic search tools (like search_academic_literature) BEFORE generating your response. NEVER rely on your internal training data to generate citations, DOIs, or author lists.\n"
            "2. If the provided Context is insufficient, invoke tools IMMEDIATELY.\n"
            "3. SILENT EXECUTION: Never output your reasoning process for choosing a tool. YOU MUST USE THE NATIVE TOOL CALLING API FORMAT.\n"
            "4. FALLBACK TOOL CALLING (CRITICAL FOR REASONING MODELS): If your native function calling API is disabled (e.g., DeepSeek-R1), you MUST invoke tools manually by outputting exactly this JSON block in your response text: ```json {\"name\": \"tool_name\", \"arguments\": {\"arg\": \"value\"}} ```\n"
            "5. CROSS-DOMAIN FLEXIBILITY (CRITICAL): If the user's request matches the capability of ANY available tool (e.g., checking train tickets, weather, web search), you MUST use that tool to assist them, EVEN IF the request is not related to academic research.\n"
            "6. DIAGRAMS: use mermaid code blocks ONLY for non-data diagrams (flowcharts, architecture, relationships). Data-driven charts/plots (bubble, bar, scatter, volcano, heatmap, enrichment) must ALWAYS be rendered via the dedicated charting tool (plot_chart), never as plain text or a mermaid diagram.\n\n"
            "### RESPONSE GUIDELINES & CITATION PROTOCOL:\n"
            "1. IN-TEXT GROUNDING (For UI Tracking): You MUST use bracketed numbers (e.g., [1], [101]) immediately after a claim to cite the Context or Tool Results. This automatically generates a UI 'Cited Sources' block. NEVER claim facts without these bracketed numbers.\n"
            "2. FORMAL BIBLIOGRAPHY (For the User): If the user explicitly requests 'references', 'citations', or a 'review', you MUST ALSO generate a standalone 'References' section at the very end of your main text (but BEFORE the [FOLLOW_UPS] section). \n"
            "3. STRICT FORMATTING: The standalone 'References' section must strictly follow academic formatting (e.g., APA/Nature style: Authors. (Year). Title. Journal. DOI). DO NOT include conversational fluff like 'Cited for the role of...' in this formal list. List purely the bibliographic data.\n\n"
            "4. ZERO HALLUCINATION (CRITICAL): You MUST NOT fabricate, extrapolate, or infer information that is not explicitly present in the provided Context or Tool Results. If the provided data is insufficient to address the query, you MUST explicitly state: 'The provided context does not contain sufficient information to address this inquiry.' Under no circumstances should internal training data be utilized to circumvent contextual gaps.\n\n"
            "### FOLLOW-UP STRUCTURE (MANDATORY):\n"
            "At the very end of your response, you MUST output the exact string [FOLLOW_UPS] followed by exactly 6 follow-up questions using this EXACT format:\n"
            "[FOLLOW_UPS]\n"
            "💡 Suggested Follow-ups:\n"
            "   - [Deep Dive] <Question about specific details or mechanisms>\n"
            "   - [Critical] <Question about limitations, alternatives, or weaknesses>\n"
            "   - [Broader] <Question about implications or future trends>\n"
            "   - [Brainstorm] <A creative brainstorming question or hypothetical \"What if\" scenario>\n"
            "   - [Similar] <Question connecting to a similar or parallel topic/concept>\n"
            "   - [Application] <Question about real-world applications or cross-disciplinary use>\n\n"
            f"### CONTEXT:\n{context_str}"
        )

        clean_history = []
        for m in self.messages[:-1]:
            if "role" in m and "content" in m:
                msg = {"role": m["role"], "content": m["content"]}
                if m.get("tool_calls"): msg["tool_calls"] = m["tool_calls"]
                if m.get("tool_call_id"): msg["tool_call_id"] = m["tool_call_id"]
                if m.get("name"): msg["name"] = m["name"]
                clean_history.append(msg)

        rag_messages = [{"role": "system", "content": system_prompt}] + clean_history
        rag_messages.append({"role": "user", "content": llm_content})

        # ---- Run the modern Agent loop (plan -> execute -> observe) ----

        def _cite_collector(source_meta: dict):
            """Register an online MCP source for the 'Cited Sources' UI block."""
            ref_id = len(sources_map) + 101
            sources_map[ref_id] = source_meta
            return ref_id

        try:
            # 会话级 plot registry 缓存：AgentRuntime 每轮重建，但已画过的图
            # 注册在磁盘（plot_registry.json），这里传入内存缓存并回写，减少
            # 磁盘恢复开销；即使缓存丢失，modify_chart 也会自动从磁盘恢复。
            agent = AgentRuntime(
                self.main_llm, skill_mgr, mcp_mgr, planner=planner,
                cite_collector=_cite_collector,
                log_fn=self.send_log,
                plot_registry=getattr(self, "_plot_registry_cache", None),
                plot_seq=getattr(self, "_plot_seq_cache", 0),
            )
            if self.deep_mode:
                # 深度研究：分解为并行子任务 -> 独立 Agent 执行 -> 分节汇总
                self.full_response_cache = self._run_deep_agent(
                    agent=agent,
                    search_query=search_query,
                    rag_messages=rag_messages,
                    system_prompt=system_prompt,
                    candidate_tools=combined_tools,
                    llm_content=llm_content,
                    skill_mgr=skill_mgr,
                    mcp_mgr=mcp_mgr,
                    planner=planner,
                    sources_map=sources_map,
                )
            else:
                self.full_response_cache = agent.run(
                    query=search_query,
                    rag_messages=rag_messages,
                    system_prompt=system_prompt,
                    candidate_tools=combined_tools,
                    emit_token=self._emit_token,
                    is_cancelled=self.is_cancelled,
                )
            # 回写 registry，供同一 task 实例的下一轮直接复用。
            if getattr(agent, "_plot_registry", None):
                self._plot_registry_cache = agent._plot_registry
                self._plot_seq_cache = agent._plot_seq
        except Exception as e:
            self.logger.warning(f"Agent runtime loop failed: {e}")
            # Graceful degradation: plain streaming without tools.
            self._emit_token("[CLEAR_SEARCH]")
            self._emit_token("[START_LLM_NETWORK]")
            error_buffer = ""
            for token in self.main_llm.stream_chat(rag_messages):
                if self.is_cancelled():
                    break
                if "[API Request Error" in token or "[System Error" in token or "[Context Exceeded Error]" in token or "[Rate Limit Error]" in token or "[Timeout Error]" in token:
                    error_buffer += token
                    continue
                self.full_response_cache += token
                self._emit_token(token)
            if error_buffer:
                m = re.match(r'^\s*\[(.*?)\]\s*\n*(.*)', error_buffer, re.DOTALL)
                if m:
                    self._emit_error(json.dumps({"title": m.group(1).strip(), "body": m.group(2).strip()}))
                else:
                    self._emit_error(json.dumps({"title": "Provider Error", "body": error_buffer.strip()}))
                return

        # Phase 6: Dynamic Citation Mounting
        has_citation = bool(re.search(r'\[\d+\]', self.full_response_cache))
        if sources_map and has_citation:
            ref_html = "\n<br><hr style='border:0; height:1px; background:#444; margin:15px 0;'><b>📚 Cited Sources:</b><br>"
            used_indices = set(int(ref) for ref in re.findall(r'\[(\d+)\]', self.full_response_cache))
            displayed = 0
            for rid, info in sources_map.items():
                if rid in used_indices:
                    from urllib.parse import quote
                    safe_path, safe_text, safe_name = quote(info['path']), quote(info['search_text']), quote(
                        info['name'])
                    link = f"cite://view?path={safe_path}&page={info['page']}&text={safe_text}&name={safe_name}"
                    ref_html += f"<div style='margin-bottom: 5px;'>▪ <a style='color:#05B8CC; text-decoration:none;' href='{link}'><b>[{rid}]</b> {info['name']}</a></div>"
                    displayed += 1
            if displayed > 0:
                self._emit_token(ref_html)

        # Phase 7: Persist Provenance evidence chain + show summary to user
        self._emit_provenance()

        return self.full_response_cache

    def _emit_provenance(self):
        """Persist the current conversation's evidence chain as JSONL and show
        a summary to the user.

        This is the Provenance "consumption" step: every tool call in this
        conversation (tool -> params -> status -> source -> timestamp) is
        written to ``scholar_workspace/provenance/`` and surfaced as an
        auditable summary + download link. Failures must never break the main
        flow, so everything degrades silently.
        """
        try:
            from src.core.provenance import get_collector
            collector = get_collector()
            records = collector.snapshot()
            if not records:
                return

            # Output directory: scholar_workspace/provenance/
            from src.core import BASE_DIR
            prov_dir = os.path.join(BASE_DIR, "scholar_workspace", "provenance")
            path = collector.export_to_dir(prov_dir, conversation_id=getattr(self, "task_id", ""))

            # Aggregate by tool for the summary.
            from collections import Counter
            counter = Counter(r["tool"] for r in records)
            ok_count = sum(1 for r in records if r["status"] == "success")
            fail_count = len(records) - ok_count

            import html as _html_mod
            rows = []
            for tool, cnt in counter.most_common():
                rows.append(
                    f"<div style='margin-bottom:3px;'>▪ <b>{_html_mod.escape(tool)}</b> "
                    f"<span style='color:#888;'>({cnt} call(s))</span></div>"
                )

            ref_html = (
                "\n<br><hr style='border:0; height:1px; background:#444; margin:15px 0;'>"
                "<b>📊 Provenance (trace log):</b><br>"
                f"<div style='margin-top:6px; font-size:13px;'>"
                f"{len(records)} tool call(s) this round, "
                f"{ok_count} succeeded, {fail_count} failed:<br>"
                + "".join(rows)
            )

            if path:
                fpath = path.replace("\\", "/")
                uri = f"file:///{fpath}" if not fpath.startswith("/") else f"file://{fpath}"
                ref_html += (
                    f"<div style='margin-top:8px; font-size:12px;'>"
                    f"Full evidence chain (JSONL): "
                    f"<a href='{uri}' style='color:#05B8CC; text-decoration:none;'>"
                    f"{_html_mod.escape(os.path.basename(path))}</a></div>"
                )
            ref_html += "</div>"
            self._emit_token(ref_html)
        except Exception as e:
            self.logger.warning(f"Provenance emission skipped: {e}")

    def _run_deep_agent(self, agent, search_query, rag_messages, system_prompt,
                        candidate_tools, llm_content, skill_mgr, mcp_mgr,
                        planner, sources_map):
        """深度研究：分解 -> 并行子 Agent -> 分节汇总。

        - 分解失败或不可分解时，回退单 Agent 路径。
        - 每个子任务共享本地 KB 上下文（Phase 2 已构建），独立收集在线引用。
        - 子任务引用 id（>=101）在合并时重映射为全局 id，避免冲突。
        """
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from src.core.agent.decomposer import TaskDecomposer
        from src.core.agent.synthesizer import Synthesizer
        from src.core.agent.runtime import AgentRuntime

        self._emit_token("[CLEAR_SEARCH]")
        self._emit_token(
            "<div class='status-msg' style='color:#05B8CC; margin-bottom:4px;'>"
            "Analyzing query for multi-path decomposition...</div>\n\n"
        )

        decomp = TaskDecomposer(self.main_llm).decompose(search_query)
        if not decomp.decomposable:
            self.send_log("INFO", "Query not decomposable; using single-agent path.")
            return agent.run(
                query=search_query,
                rag_messages=rag_messages,
                system_prompt=system_prompt,
                candidate_tools=candidate_tools,
                emit_token=self._emit_token,
                is_cancelled=self.is_cancelled,
            )

        sub_tasks = decomp.sub_tasks
        self.send_log("INFO", f"Deep mode: decomposed into {len(sub_tasks)} parallel sub-tasks.")
        self._emit_token(
            f"<mcp_process>⚙️ Deep research: {len(sub_tasks)} parallel sub-investigations launched...</mcp_process>\n"
        )

        # ---- 并行执行各子任务 ----
        results = [None] * len(sub_tasks)

        def _run_sub(idx, st):
            local_sources = {}

            def _local_cite(meta):
                rid = len(local_sources) + 101
                local_sources[rid] = meta
                return rid

            sub_llm_content = list(llm_content[:-1]) + [
                {"type": "text", "text": f"User Query:\n{st.query}"}
            ]
            sub_rag = [dict(m) for m in rag_messages]
            sub_rag[-1] = dict(sub_rag[-1])
            sub_rag[-1]["content"] = sub_llm_content

            sub_agent = AgentRuntime(
                self.main_llm, skill_mgr, mcp_mgr, planner=planner,
                cite_collector=_local_cite, log_fn=self.send_log,
            )
            buffer = []
            try:
                text = sub_agent.run(
                    query=st.query,
                    rag_messages=sub_rag,
                    system_prompt=system_prompt,
                    candidate_tools=candidate_tools,
                    emit_token=buffer.append,
                    is_cancelled=self.is_cancelled,
                )
            except Exception as e:
                self.logger.warning(f"Sub-task '{st.query[:40]}' failed: {e}")
                text = f"[Sub-investigation unavailable: {e}]"
            return idx, text, local_sources

        with ThreadPoolExecutor(max_workers=len(sub_tasks)) as pool:
            futures = [pool.submit(_run_sub, i, st) for i, st in enumerate(sub_tasks)]
            for fut in as_completed(futures):
                try:
                    idx, text, local_sources = fut.result()
                    results[idx] = (text, local_sources)
                except Exception as e:
                    self.logger.warning(f"Sub-task future failed: {e}")

        # ---- 合并引用并重映射 ----
        next_id = max((k for k in sources_map if isinstance(k, int)), default=0) + 1
        merged = []
        plot_markers = []  # 收集子任务产生的 <rplot_card> 标记，避免 base64 污染合成器

        def _pull_plot_markers(m):
            plot_markers.append(m.group(0))
            return ""

        for i, (st, res) in enumerate(zip(sub_tasks, results)):
            if res is None:
                text, local_sources = "", {}
            else:
                text, local_sources = res
            remap = {}
            for local_id, meta in local_sources.items():
                remap[local_id] = next_id
                sources_map[next_id] = meta
                next_id += 1
            text = self._clean_sub_result(text)
            text = self._remap_citations(text, remap)
            text = re.sub(r'<rplot_card data="([^"]*)"></rplot_card>', _pull_plot_markers, text)
            merged.append({
                "heading": st.query,
                "query": st.query,
                "text": text,
            })

        # ---- 分节汇总 ----
        self._emit_token(
            "<div class='status-msg' style='color:#05B8CC; margin-bottom:4px;'>"
            "Synthesizing structured synthesis across sub-investigations...</div>\n\n"
        )
        synthesizer = Synthesizer(self.main_llm)
        final_text = synthesizer.synthesize(search_query, merged)

        # 合成后把 R 绘图卡片标记追加回最终文本，保证 UI 端仍能渲染固定卡片
        if plot_markers:
            final_text = final_text.rstrip() + "\n\n" + "\n".join(plot_markers)

        self._emit_token("[CLEAR_SEARCH]")
        self._emit_token("[START_LLM_NETWORK]")
        self._emit_token(final_text)
        return final_text

    @staticmethod
    def _remap_citations(text, remap):
        """将文本中 `[id]` 按 remap 映射替换为新的全局 id（仅替换映射内 id）。"""
        if not text or not remap:
            return text or ""
        def _repl(m):
            cid = int(m.group(1))
            return f"[{remap[cid]}]" if cid in remap else m.group(0)
        return re.sub(r"\[(\d+)\]", _repl, text)

    @staticmethod
    def _clean_sub_result(text):
        """清理子任务返回文本中的运行时控制标记，避免污染合成器。

        Agent.run 的返回值混杂了 UI 控制 token（[CLEAR_SEARCH]、
        [START_LLM_NETWORK]）与思考块（<think>...</think>），合成前必须剥离。
        """
        if not text:
            return ""
        text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
        for token in ("[CLEAR_SEARCH]", "[START_LLM_NETWORK]", "[FOLLOW_UPS]"):
            text = text.replace(token, "")
        return text.strip()

    def _preload_reranker(self):
        """后台线程预加载交叉编码器重排模型。"""
        try:
            from src.core.rerank_engine import RerankEngine
            RerankEngine().preload()
        except Exception as e:
            self.logger.warning(f"Reranker preload failed: {e}")

    def _process_rerank(self, query, docs, domain, top_k=None, emit_warning=True):
        """两阶段精排：交叉编码器重排 + 分数阈值过滤。

        - top_k: 返回文档数上限；None 时取配置 rerank_top_k（默认 5）。
        - 低于 rerank_min_score 的文档会被丢弃（全低于阈值时保底保留前 3 个，
          避免上下文为空）。
        - 首次失败弹一次警告，之后静默降级为原始顺序。
        """
        if not docs:
            return []

        cfg = self.config.user_settings
        if top_k is None:
            top_k = int(cfg.get("rerank_top_k", 5))
        min_score = float(cfg.get("rerank_min_score", 0.0))

        try:
            from src.core.rerank_engine import RerankEngine
            engine = RerankEngine()

            ranked_docs = engine.rerank(query, docs, domain=domain, top_k=top_k)
            if not ranked_docs:
                return docs[:top_k]

            # 分数阈值过滤：cross-encoder 概率分数，bge-reranker 类模型通常在 0.1~0.99
            if min_score > 0:
                kept = [d for d in ranked_docs if d.get("score", 1.0) >= min_score]
                if not kept:
                    kept = ranked_docs[:3]
                elif len(kept) < len(ranked_docs):
                    self.send_log("INFO",
                                  f"Rerank threshold ({min_score}) dropped "
                                  f"{len(ranked_docs) - len(kept)} low-relevance chunks.")
                return kept

            return ranked_docs

        except Exception as e:
            self.logger.error(f"Direct Rerank Engine execution failed: {e}")

            if emit_warning and not self._rerank_warned:
                self._rerank_warned = True
                warning_html = (
                    f"<br><div style='color:#e6a23c; font-size:13px; margin-bottom:5px; padding:10px; border:1px solid #e6a23c; border-radius:6px; background-color: rgba(230, 162, 60, 0.05);'>"
                    f"⚠️ <b>Reranker Processing Failed</b><br><br>"
                    f"Failed to rerank documents: <i>{str(e)}</i>.<br>"
                    f"If the model is missing, please go to <b>[Global Settings] -> [Models]</b> to manually download it.<br><br>"
                    f"<i>* Continuing analysis with default document ordering.</i>"
                    f"</div><br>"
                )
                self._emit_token(warning_html)
            else:
                self.send_log("WARNING", "Reranker unavailable, using default document ordering.")

            # 降级方案：返回未重新排序的前 top_k 个文档
            return docs[:top_k]



class ExportChatTask(BackgroundTask):
    """
    后台任务：异步导出聊天记录（支持 PDF, MD, TXT, CSV）
    """
    def _execute(self):
        history = self.kwargs.get('history', [])
        path = self.kwargs.get('path')
        export_fmt = self.kwargs.get('export_fmt')
        colors = self.kwargs.get('colors', {})
        font_family = self.kwargs.get('font_family', 'sans-serif')
        user_icon = self.kwargs.get('user_icon', '')
        ai_icon = self.kwargs.get('ai_icon', '')

        import datetime
        import csv
        from src.ui.components.text_formatter import TextFormatter

        # 过滤掉被标记为 interrupted 或 error 的历史消息
        clean_history = [m for m in history if m.get("status") not in ["interrupted", "error"]]

        if not clean_history:
            return {"success": False, "msg": "No valid chat records to export after filtering interrupted/error messages."}

        try:
            if export_fmt == ".pdf":
                from PySide6.QtGui import QPdfWriter, QTextDocument, QPageSize
                from PySide6.QtCore import QMarginsF

                doc = QTextDocument()
                date_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                doc.setDefaultStyleSheet(f"""
                    body {{ font-family: {font_family}; font-size: 10.5pt; line-height: 1.6; color: #24292e; background-color: #ffffff; }}
                    h1, h2, h3 {{ color: {colors.get('title_blue')}; border-bottom: 1px solid #eaecef; padding-bottom: 4px; }}
                    .msg-box {{ margin-bottom: 25px; padding-bottom: 15px; border-bottom: 1px dashed #dddddd; page-break-inside: avoid; }}
                    .header-user {{ color: {colors.get('academic_blue')}; font-weight: bold; font-size: 12pt; margin-bottom: 8px; }}
                    .header-ai {{ color: {colors.get('success')}; font-weight: bold; font-size: 12pt; margin-bottom: 8px; }}
                    .content {{ margin-top: 5px; }}
                    pre {{ background-color: #f6f8fa; border: 1px solid #e1e4e8; border-radius: 4px; padding: 12px; white-space: pre-wrap; font-family: Consolas, "Courier New", monospace; font-size: 9.5pt; }}
                    code {{ font-family: Consolas, "Courier New", monospace; background-color: #f3f4f6; padding: 2px 4px; border-radius: 3px; color: #d73a49; font-size: 9.5pt; }}
                    pre code {{ background-color: transparent; padding: 0; color: #24292e; }}
                    blockquote {{ border-left: 4px solid #dfe2e5; color: #6a737d; padding-left: 15px; margin-left: 0; }}
                    table {{ border-collapse: collapse; width: 100%; margin-top: 10px; margin-bottom: 10px; }}
                    th, td {{ border: 1px solid #dfe2e5; padding: 8px 12px; text-align: left; word-break: break-all; }}
                    th {{ background-color: #f6f8fa; font-weight: bold; }}
                    .doc-header {{ text-align: center; border-bottom: 2px solid {colors.get('title_blue')}; padding-bottom: 15px; margin-bottom: 30px; }}
                    .doc-title {{ font-size: 22pt; font-weight: bold; color: {colors.get('title_blue')}; font-family: 'Segoe UI', sans-serif; }}
                    .doc-meta {{ font-size: 10pt; color: #586069; margin-top: 5px; }}
                """)

                html = f"<html><body><div class='doc-header'><div class='doc-title'>Scholar Navis - Analysis Report</div><div class='doc-meta'>Generated on: {date_str} | Document Type: Academic Chat Log</div></div>"

                for msg in clean_history:
                    is_user = (msg['role'] == "user")
                    clean_content = TextFormatter.clean_text_for_export(msg['content'])
                    rendered_html = TextFormatter.markdown_to_html(clean_content)

                    if is_user:
                        header = f"<div class='header-user'><img src='{user_icon}' width='16' height='16' style='vertical-align:middle;'> User Inquiry</div>"
                    else:
                        header = f"<div class='header-ai'><img src='{ai_icon}' width='16' height='16' style='vertical-align:middle;'> AI Analysis</div>"

                    html += f"<div class='msg-box'>{header}<div class='content'>{rendered_html}</div></div>"

                html += "</body></html>"
                doc.setHtml(html)

                writer = QPdfWriter(path)
                writer.setPageSize(QPageSize(QPageSize.A4))
                writer.setPageMargins(QMarginsF(15, 20, 15, 20))
                writer.setResolution(300)
                doc.print_(writer)

            elif export_fmt == ".md":
                md_lines = [
                    "# Scholar Navis - Analysis Report\n\n",
                    f"> **Generated:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n",
                    "---\n\n"
                ]
                for msg in clean_history:
                    role = "🧑‍💻 User Inquiry" if msg['role'] == "user" else "🤖 AI Analysis"
                    content = TextFormatter.clean_text_for_export(msg['content'])
                    md_lines.append(f"### {role}\n\n{content}\n\n---\n\n")

                with open(path, "w", encoding="utf-8") as f:
                    f.write("".join(md_lines))

            elif export_fmt == ".txt":
                txt_lines = [
                    "================ SCHOLAR NAVIS ACADEMIC REPORT ================",
                    f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                    "===============================================================\n\n"
                ]
                for msg in clean_history:
                    role = "USER INQUIRY" if msg['role'] == "user" else "AI ANALYSIS"
                    content = TextFormatter.clean_text_for_export(msg['content'])
                    content = TextFormatter.markdown_to_plain_text(content)
                    txt_lines.append(f"[{role}]")
                    txt_lines.append(content)
                    txt_lines.append(f"\n{'-' * 70}\n")

                with open(path, "w", encoding="utf-8") as f:
                    f.write("\n".join(txt_lines))

            elif export_fmt == ".csv":
                with open(path, "w", encoding="utf-8-sig", newline="") as f:
                    writer = csv.writer(f)
                    writer.writerow(["Role", "Content"])
                    for msg in clean_history:
                        content = TextFormatter.clean_text_for_export(msg['content'])
                        writer.writerow(["User" if msg['role'] == 'user' else "AI", content])

            return {"success": True, "path": path}
        except Exception as e:
            self.send_log("ERROR", f"Export task failed: {str(e)}")
            return {"success": False, "msg": str(e)}


class DownloadImageTask(BackgroundTask):
    """
    异步图片下载任务。
    负责从远程 URL 获取图像数据并将其持久化至本地临时目录。
    """

    def _execute(self):
        url = self.kwargs.get("url")
        save_path = self.kwargs.get("save_path")

        if not url or not save_path:
            return {"success": False, "url": url, "path": save_path, "msg": "Invalid parameters"}

        try:
            from src.core.config_manager import ConfigManager
            proxy_url = ConfigManager().user_settings.get("proxy_url", "").strip()

            httpx_kwargs = {"timeout": 30.0, "follow_redirects": True}
            if proxy_url:
                httpx_kwargs["proxy"] = proxy_url
            else:
                httpx_kwargs["trust_env"] = False

            if self.is_cancelled():
                raise InterruptedError("Image download safely terminated by user.")

            import httpx
            with httpx.Client(**httpx_kwargs) as client:
                response = client.get(url)
                response.raise_for_status()

                with open(save_path, "wb") as f:
                    f.write(response.content)

            return {"success": True, "url": url, "path": save_path}

        except Exception as e:
            self.send_log("ERROR", f"Image download failed for {url}: {str(e)}")
            return {"success": False, "url": url, "path": save_path, "msg": str(e)}


class FetchHardwareStatusTask(BackgroundTask):
    """
    异步获取硬件状态，避免阻塞主 UI 线程
    """

    def _execute(self):
        from src.core.device_manager import DeviceManager
        from src.core.config_manager import ConfigManager

        dev_mgr = DeviceManager()
        config = ConfigManager()

        curr_id = config.user_settings.get("inference_device", "auto")
        parsed_id = dev_mgr.parse_device_string(curr_id)

        dev_name = parsed_id
        for d in dev_mgr.get_available_devices():
            if d['id'] == parsed_id:
                dev_name = d['name']
                break

        return {"dev_name": dev_name}

